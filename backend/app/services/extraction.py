"""
Text extraction service.
Handles PDF (PyPDF2), DOCX (python-docx), and TXT files.
Includes contract validation.
"""
import re
from typing import Tuple

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None


# Keywords that indicate a legal contract
CONTRACT_KEYWORDS = [
    "agreement", "party", "parties", "terms", "conditions",
    "liability", "indemnity", "termination", "confidential",
    "obligations", "shall", "hereby", "whereas", "covenant",
    "warrant", "represent", "governing law", "jurisdiction",
    "breach", "remedy", "clause", "section", "article",
    "execute", "binding", "contract", "license", "lessee",
    "lessor", "tenant", "landlord", "employer", "employee",
]

MIN_KEYWORD_MATCHES = 3  # Require at least 3 distinct keyword hits


def extract_text(file_path: str, file_type: str) -> Tuple[str, int]:
    """
    Extract text from a file.
    Returns: (text, page_count)
    """
    file_type = file_type.lower().strip(".")

    if file_type == "pdf":
        if not PdfReader:
            raise RuntimeError("PDF extraction module (PyPDF2) is not installed on the server.")
        return _extract_pdf(file_path)
    elif file_type in ("docx", "doc"):
        if not Document:
            raise RuntimeError("DOCX extraction module (python-docx) is not installed on the server.")
        return _extract_docx(file_path)
    elif file_type == "txt":
        return _extract_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def validate_contract(text: str) -> bool:
    """
    Check if the extracted text looks like a legal contract.
    Returns True if it contains enough legal keywords.
    """
    if not text or len(text.strip()) < 100:
        return False

    text_lower = text.lower()
    matches = sum(1 for kw in CONTRACT_KEYWORDS if kw in text_lower)
    return matches >= MIN_KEYWORD_MATCHES


def _extract_pdf(file_path: str) -> Tuple[str, int]:
    """Extract text from PDF using PyPDF2."""
    try:
        reader = PdfReader(file_path)
        page_count = len(reader.pages)
        extracted_pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text and text.strip():
                extracted_pages.append(text)
        return "\n\n".join(extracted_pages), page_count
    except Exception as e:
        raise RuntimeError(f"PDF extraction failed: {str(e)}")


def _extract_docx(file_path: str) -> Tuple[str, int]:
    """Extract text from DOCX using python-docx."""
    try:
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table content
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n\n".join(paragraphs), 1
    except Exception as e:
        raise RuntimeError(f"DOCX extraction failed: {str(e)}")


def _extract_txt(file_path: str) -> Tuple[str, int]:
    """Extract text from plain text file."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        return text, 1
    except Exception as e:
        raise RuntimeError(f"TXT extraction failed: {str(e)}")
